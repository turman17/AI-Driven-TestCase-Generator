from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
from upload_file import upload_file
from docx_reader import read_docx_content_in_order
from openai import AzureOpenAI
from config import azure_endpoint1, api_version1, parse_instructions, output_instructions
from key import api_key1
import os
import json
import csv
from docx import Document

app = Flask(__name__, static_url_path='/static', static_folder='results')
CORS(app)

# loadAzureOpenAI
# This function initializes and returns an instance of the AzureOpenAI client.
# It uses predefined variables for the Azure endpoint, API key, and API version.
# return: An instance of AzureOpenAI client.
def loadAzureOpenAI():
    client = AzureOpenAI(
    azure_endpoint=azure_endpoint1, 
    api_key=api_key1,  
    api_version=api_version1
)
    return client

# open_ai
# This function uses the provided client to interact with OpenAI's GPT-4 model.
# It sends a series of messages to the model and receives a completion.
# The messages consist of a system message and a user message, which is the input text.
# The function configures the model's behavior with parameters like temperature, max_tokens, top_p, frequency_penalty, and presence_penalty.
# 
# Parameters:
# client: The OpenAI client to use for the interaction.
# text: The user message to send to the model.
# 
# return: The content of the first choice in the completion.
def open_ai(client, text):
    messages = [
        {"role": "system", "content": "Your system message goes here."},
        {"role": "user", "content": text}
    ]
    completion = client.chat.completions.create(
        model="gpt4",
        messages=messages,
        temperature=0.73,
        max_tokens=4096,
        top_p=0.90,
        frequency_penalty=0,
        presence_penalty=0,
        stop=None
    )

    return completion.choices[0].message.content

# process_file
# This route is the main entry point for processing a file with OpenAI.
# It first uploads the file from the request and checks if the upload was successful.
# If the upload was successful, it reads the content of the file in order.
# It then initializes the Azure OpenAI client and sends the file content to OpenAI's GPT-4 model.
# If there are any extra instructions in the request, they are added to the beginning of the output.
# Finally, it sends the output to the OpenAI model again and returns the result.
# 
# Parameters:
# request: The HTTP request containing the file and any extra instructions.
# 
# return: The output from the OpenAI model and a 200 status code.
@app.route('/', methods=['POST'])
def process_file():
    file_path = upload_file(request)
    if not isinstance(file_path, str):
        return file_path
    text = read_docx_content_in_order(file_path)
    client = loadAzureOpenAI()
    combined_text = parse_instructions + text
    tokens = open_ai(client, combined_text)
    extraInstructions = request.form['extraInstructions']
    if (extraInstructions):
        tokens = output_instructions + extraInstructions + tokens
    else:
        tokens = output_instructions + tokens
    message_output = open_ai(client, tokens)
    
    return message_output, 200


# json_to_docx
# This function converts a JSON string into a Word document using the python-docx library.
# The JSON string should contain a title and a list of test cases.
# Each test case should have an ID, title, description, pre-conditions, requirement, and actions.
# Each action should have a step and expected results.
# The function creates a new Word document, adds a heading for the title, and then adds a section for each test case.
# Each test case section includes a heading, paragraphs for the ID, title, description, pre-conditions, and requirement, and a table for the actions.
# The table has two columns: one for the action step and one for the expected results.
# If there are more test cases, the function adds a page break before the next one.
# 
# Parameters:
# json_str: The JSON string to convert into a Word document.
# 
# return: The created Word document.
def json_to_docx(json_str):
    data = json.loads(json_str)
    doc = Document()
    doc.add_heading(data['title'], level=1)
    
    for index, test_case in enumerate(data['testCases']):
        doc.add_heading('Test Case ' + test_case['testCaseId'], level=2)
        doc.add_paragraph('Test Case\n' + test_case['testCaseId'])
        doc.add_paragraph('Title\n' + test_case['title'])
        doc.add_paragraph('Description\n' + test_case['description'])
        doc.add_paragraph('Pre-Conditions\n' + '\n'.join(test_case['preConditions']))
        doc.add_paragraph('Requirement\n' + ', '.join(test_case['requirement']))
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Action'
        hdr_cells[1].text = 'Expected Result'
        
        for action in test_case['actions']:
            row = table.add_row().cells
            row[0].text = action['step']
            row[1].text = '\nAND\n'.join(action['expectedResults'])
        
        if index < len(data['testCases']) - 1:
            doc.add_page_break()

    return doc


# json_to_csv
# This function converts a JSON string into a CSV file.
# The JSON string should contain a title and a list of test cases.
# Each test case should have a title and a list of actions.
# Each action should have a step and expected results.
# The function creates a new CSV file with headers for Id, Work Item Type, Title, Test Step, Step Action, Step Expected, Area Path, Assigned To, and State.
# It then adds a row for each test case and a row for each action in the test case.
# The test case row includes the test case title and some default values.
# The action row includes the action step and expected results.
# The CSV file is saved in the 'results' directory with the name 'Test cases.csv'.
# 
# Parameters:
# json_str: The JSON string to convert into a CSV file.
# 
# return: None. The function writes the CSV file to disk.
def json_to_csv(json_str):
    data = json_str
    csv_header = ['Id', 'Work Item Type', 'Title', 'Test Step', 'Step Action', 'Step Expected', 'Area Path', 'Assigned To', 'State']
    csv_rows = []
    work_item_title = data['title']

    # Process each test case
    for testCase in data['testCases']:
        # Work Item row
        csv_rows.append([
            "",  # Id is blank
            "Test Case",  # Work Item Type
            f"{work_item_title} - {testCase['title']}",  # Title
            "",  # Test Step
            "",  # Step Action
            "",  # Step Expected
            "OBSERVAERO",  # Area Path
            "Ricardo1 SANTOS <ricardo1.santos@vinci-energies.net>",  # Assigned To
            "Design"  # State
        ])

        for index, action in enumerate(testCase['actions']):
            expected_results_combined = "\n".join(action['expectedResults'])
            
            csv_rows.append([
                "",  # Id
                "",  # Work Item Type
                "",  # Title
                str(index + 1),  # Test Step
                action['step'],  # Step Action
                expected_results_combined,  # Step Expected (combined)
                "",  # Area Path
                "",  # Assigned To
                ""  # State
            ])

    filename = 'Test cases.csv'
    file_path = os.path.join('results', filename)

    with open(file_path, mode='w', newline='') as file:
        writer = csv.writer(file)
        writer.writerow(csv_header)
        writer.writerows(csv_rows)

# generate_file
# This route is the main entry point for generating a Word document and a CSV file from a JSON string.
# It first receives the JSON string from the request and checks if the data was received.
# If the data was not received, it returns an error.
# If the 'results' directory does not exist, it creates it.
# It then calls the json_to_docx function and the json_to_csv function to convert the JSON string into a Word document and a CSV file.
# It defines the filename as 'Test cases' and saves the Word document in the 'results' directory.
# Finally, it returns the filename and a 200 status code.
# 
# Parameters:
# request: The HTTP request containing the JSON string.
# 
# return: The filename and a 200 status code.
@app.route('/generate', methods=['POST'])
def generate_file():
    data = request.get_json()
    if not data:
        return jsonify({'error': 'No data provided'}), 400

    if not os.path.exists('results'):
        os.makedirs('results')

    json_str = json.dumps(data)
    doc = json_to_docx(json_str)
    json_to_csv(data)

    filename = 'Test cases'
    filenameDocx = filename + '.docx'

    file_path = os.path.join('results', filenameDocx)

    doc.save(file_path)

    return jsonify({'file_path': filename}), 200


# check_connection
# This route is used to check if the Flask server is running.
# It responds to GET requests at the root URL ("/").
# When a GET request is received, it returns a JSON object with a message saying "Flask server is running!".
# This can be used to verify that the server is up and able to respond to requests.
# 
# Parameters:
# None.
# 
# return: A JSON object with a message saying "Flask server is running!".
@app.route('/', methods=['GET'])
def check_connection():
    return jsonify({'message': 'Flask server is running!'})


# serve_static
# This route is used to serve static files from the 'results' directory.
# It responds to GET requests at the "/static/<path:path>" URL.
# When a GET request is received, it returns the file at the specified path in the 'results' directory.
# This can be used to download the files generated by the application.
# 
# Parameters:
# path: The path to the file in the 'results' directory.
# 
# return: The requested file from the 'results' directory.
@app.route('/static/<path:path>')
def serve_static(path):
    return send_from_directory('results', path)


# Main entry point
# This is the main entry point of the Flask application.
# If this script is run directly (not imported as a module), it starts the Flask development server.
# The server is set to listen on all IP addresses (0.0.0.0) and port 8000.
# The debug mode is turned on, which means that the server will automatically reload if code changes are detected, and it will provide detailed error messages in case of an error.
# 
# Parameters:
# None.
# 
# return: None. This function starts the Flask server and does not return.
if __name__ == '__main__':
    app.run('0.0.0.0', 8000, debug=True)